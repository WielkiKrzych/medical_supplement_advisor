from docx import Document
import os

file_path = "/Users/wielkikrzych/medical-supplement-advisor/examples/sample_blood_tests.docx"
doc = Document(file_path)

print(f"--- Full Analysis of {file_path} ---")

print("\n[Paragraphs]")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        print(f"[{i}] {text}")

for t_idx, table in enumerate(doc.tables):
    print(f"\n[Table {t_idx}]")
    for r_idx, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        print(f"  Row {r_idx}: {cells}")

