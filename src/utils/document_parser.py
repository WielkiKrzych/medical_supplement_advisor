"""
Document parser for extracting blood test results from PDF and DOCX files.

This module provides functionality to parse blood test documents and convert
them to the JSON format required by the Medical Supplement Advisor.
"""

import re
import io
from pathlib import Path
from typing import Dict, List, Optional, Union
from dataclasses import dataclass

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pdfplumber

    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import fitz

    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image

    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

from src.utils.exceptions import DataLoaderError


@dataclass
class PatientData:
    """Patient information extracted from document."""

    name: str
    surname: str
    age: int
    conditions: List[str]

    def to_dict(self) -> Dict:
        """Convert to dictionary format expected by JSON parser."""
        return {
            "name": self.name,
            "surname": self.surname,
            "age": self.age,
            "conditions": self.conditions,
        }


@dataclass
class BloodTest:
    """Single blood test result."""

    name: str
    value: float
    unit: str

    def to_dict(self) -> Dict:
        """Convert to dictionary format expected by JSON parser."""
        return {"name": self.name, "value": self.value, "unit": self.unit}


class DocumentParser:
    """
    Parser for blood test documents in PDF and DOCX formats.

    Supports automatic format detection and extraction of patient data
    and blood test results.
    """

    def __init__(self):
        """Initialize the document parser."""
        if not DOCX_AVAILABLE and not PDFPLUMBER_AVAILABLE:
            raise DataLoaderError(
                "No document parsing libraries available. "
                "Install pdfplumber and/or python-docx."
            )

    def parse_document(self, file_path: Union[str, Path]) -> Dict:
        """
        Parse a blood test document (PDF or DOCX).

        Args:
            file_path: Path to the document file

        Returns:
            Dictionary with 'patient' and 'blood_tests' keys

        Raises:
            DataLoaderError: If file cannot be parsed
        """
        file_path = Path(file_path)

        if not file_path.exists():
            raise DataLoaderError(f"File not found: {file_path}")

        # Detect format and parse accordingly
        if file_path.suffix.lower() == ".docx":
            if not DOCX_AVAILABLE:
                raise DataLoaderError(
                    "python-docx library not available. "
                    "Install it with: pip install python-docx"
                )
            return self._parse_docx(file_path)

        elif file_path.suffix.lower() == ".pdf":
            if not PDFPLUMBER_AVAILABLE:
                raise DataLoaderError(
                    "pdfplumber library not available. "
                    "Install it with: pip install pdfplumber"
                )
            result = self._parse_pdf(file_path)

        else:
            raise DataLoaderError(
                f"Unsupported file format: {file_path.suffix}. "
                "Supported formats: .pdf, .docx"
            )

        if not result["patient"]["name"] and not result["patient"]["surname"]:
            result["patient"]["name"] = "Pacjent"
            result["patient"]["surname"] = "Z pliku PDF"

        if not result["patient"]["surname"]:
            result["patient"]["surname"] = " "

        if result["patient"]["age"] == 0:
            result["patient"]["age"] = 35

        return result

    def _parse_docx(self, file_path: Path) -> Dict:
        """
        Parse DOCX file and extract patient data and blood tests.

        Expected structure:
        - Table 0: Patient information (2 rows x 4 columns)
        - Table 1: Blood test results (header + data rows, 3 columns)

        Args:
            file_path: Path to DOCX file

        Returns:
            Dictionary with patient and blood_tests data
        """
        try:
            doc = Document(str(file_path))

            if len(doc.tables) < 2:
                raise DataLoaderError(
                    f"Invalid DOCX format: Expected at least 2 tables, found {len(doc.tables)}"
                )

            # Parse patient data from Table 0
            patient_table = doc.tables[0]
            patient_data = self._extract_patient_from_table(patient_table)

            # Parse blood tests from Table 1
            tests_table = doc.tables[1]
            blood_tests = self._extract_blood_tests_from_table(tests_table)

            return {
                "patient": patient_data.to_dict(),
                "blood_tests": [test.to_dict() for test in blood_tests],
            }

        except Exception as e:
            raise DataLoaderError(f"Failed to parse DOCX file: {str(e)}")

    def _parse_pdf(self, file_path: Path) -> Dict:
        """
        Parse PDF file and extract patient data and blood tests.

        This method attempts multiple strategies:
        1. Extract tables from PDF
        2. Extract from text using pdfplumber
        3. Fallback to OCR if text is garbled

        Args:
            file_path: Path to PDF file

        Returns:
            Dictionary with patient and blood_tests data
        """
        try:
            with pdfplumber.open(str(file_path)) as pdf:
                all_tables = []
                for page in pdf.pages:
                    tables = page.extract_tables()
                    if tables:
                        all_tables.extend(tables)

                if len(all_tables) >= 2:
                    patient_data = self._extract_patient_from_table(all_tables[0])
                    blood_tests = self._extract_blood_tests_from_table(all_tables[1])
                else:
                    patient_data, blood_tests = self._extract_from_text(pdf)

                # Check if extracted text is garbled (CID encoding)
                combined_text = ""
                for page in pdf.pages:
                    combined_text += page.extract_text() + "\n"

                if self._is_text_garbled(combined_text):
                    return self._parse_pdf_with_ocr(file_path)

                return {
                    "patient": patient_data.to_dict(),
                    "blood_tests": [test.to_dict() for test in blood_tests],
                }

        except Exception as e:
            try:
                return self._parse_pdf_with_ocr(file_path)
            except Exception as ocr_error:
                raise DataLoaderError(
                    f"Failed to parse PDF file. Original error: {str(e)}. OCR also failed: {str(ocr_error)}"
                )

    def _extract_patient_from_table(self, table) -> PatientData:
        """
        Extract patient data from a table row.

        Expected format:
        Row 0: [Imię, Nazwisko, Wiek, Schorzenia]
        Row 1: [Jan, Nowak, 42, osteoporoza]

        Args:
            table: python-docx Table object or pdfplumber table

        Returns:
            PatientData object
        """
        try:
            # Handle both docx and pdfplumber table formats
            if hasattr(table, "rows"):  # docx format
                # Extract header and data rows
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                data = [cell.text.strip() for cell in table.rows[1].cells]
            else:  # pdfplumber format (list of lists)
                headers = [str(cell).strip() if cell else "" for cell in table[0]]
                data = [str(cell).strip() if cell else "" for cell in table[1]]

            # Map columns to fields
            field_map = {}
            for i, header in enumerate(headers):
                if i < len(data):
                    field_map[header] = data[i]

            # Extract name
            name = field_map.get("Imię", field_map.get("Imie", ""))
            surname = field_map.get("Nazwisko", "")

            # Extract age
            age_str = field_map.get("Wiek", "0")
            try:
                age = int(re.sub(r"[^0-9]", "", age_str))
            except ValueError:
                age = 0

            # Extract conditions
            conditions_str = field_map.get("Schorzenia", "")
            conditions = []
            if conditions_str:
                # Split by comma or semicolon
                conditions = [
                    c.strip() for c in re.split(r"[,;]", conditions_str) if c.strip()
                ]

            return PatientData(
                name=name, surname=surname, age=age, conditions=conditions
            )

        except Exception as e:
            raise DataLoaderError(f"Failed to extract patient data: {str(e)}")

    def _extract_blood_tests_from_table(self, table) -> List[BloodTest]:
        """
        Extract blood test results from a table.

        Expected format:
        Row 0: [Badanie, Wartość, Jednostka] (headers)
        Row 1+: [Test Name, Value, Unit]

        Args:
            table: python-docx Table object or pdfplumber table

        Returns:
            List of BloodTest objects
        """
        try:
            blood_tests = []

            # Handle both docx and pdfplumber table formats
            if hasattr(table, "rows"):  # docx format
                rows = table.rows[1:]  # Skip header row
                for row in rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) >= 3 and cells[0]:  # Skip empty rows
                        test = self._parse_blood_test_row(cells)
                        if test:
                            blood_tests.append(test)
            else:  # pdfplumber format
                rows = table[1:]  # Skip header row
                for row in rows:
                    cells = [str(cell).strip() if cell else "" for cell in row]
                    if len(cells) >= 3 and cells[0]:  # Skip empty rows
                        test = self._parse_blood_test_row(cells)
                        if test:
                            blood_tests.append(test)

            return blood_tests

        except Exception as e:
            raise DataLoaderError(f"Failed to extract blood tests: {str(e)}")

    def _parse_blood_test_row(self, cells: List[str]) -> Optional[BloodTest]:
        """
        Parse a single blood test row.

        Args:
            cells: List of cell values [name, value, unit]

        Returns:
            BloodTest object or None if parsing fails
        """
        try:
            name = cells[0].strip()
            value_str = cells[1].strip()
            unit = cells[2].strip() if len(cells) > 2 else ""

            if not name or not value_str:
                return None

            # Clean up value string and convert to float
            value_str = re.sub(r"[^\d.,-]", "", value_str)
            value_str = value_str.replace(",", ".")  # Handle European format
            value = float(value_str)

            return BloodTest(name=name, value=value, unit=unit)

        except (ValueError, IndexError):
            # Skip rows that cannot be parsed
            return None

    def _extract_from_text(self, pdf) -> tuple:
        """
        Fallback method: Extract data from PDF text content.

        This uses regex patterns to find patient data and blood test results
        in plain text format.

        Args:
            pdf: pdfplumber PDF object

        Returns:
            Tuple of (PatientData, List[BloodTest])
        """
        # Extract all text from PDF
        full_text = ""
        for page in pdf.pages:
            full_text += page.extract_text() + "\n"

        # Extract patient information using regex
        # Pattern: "Imię: Jan Nazwisko: Nowak Wiek: 42"
        name_match = re.search(
            r"[Ii]mi[ęe][:\s]+([A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+)", full_text
        )
        surname_match = re.search(
            r"[Nn]azwisko[:\s]+([A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+)", full_text
        )
        age_match = re.search(r"[Ww]iek[:\s]+(\d+)", full_text)

        patient_data = PatientData(
            name=name_match.group(1) if name_match else "",
            surname=surname_match.group(1) if surname_match else "",
            age=int(age_match.group(1)) if age_match else 0,
            conditions=[],
        )

        # Extract blood tests using regex
        # Pattern: "Vitamina D: 22 ng/mL"
        blood_tests = []
        test_pattern = (
            r"([A-ZĄĆĘŁŃÓŚŹŻ][a-zA-Ząćęłńóśźż\s()]+?)[:\s]+([\d.,]+)\s*([a-zA-Z/%\s]+)"
        )

        for match in re.finditer(test_pattern, full_text):
            try:
                name = match.group(1).strip()
                value_str = match.group(2).replace(",", ".")
                value = float(value_str)
                unit = match.group(3).strip()

                blood_tests.append(BloodTest(name=name, value=value, unit=unit))
            except (ValueError, IndexError):
                continue

        return patient_data, blood_tests

    def _is_text_garbled(self, text: str) -> bool:
        """
        Check if extracted text is garbled (contains CID codes).

        CID encoding appears as many (cid:XXX) patterns and indicates
        the PDF uses special character encoding that needs OCR.

        Args:
            text: Extracted text to check

        Returns:
            True if text appears garbled
        """
        if not text:
            return True

        # Count CID patterns - more than 10 suggests CID encoding
        cid_count = text.count("(cid:")
        if cid_count > 10:
            return True

        # Check if text has very few readable characters
        readable_chars = sum(1 for c in text if c.isalnum() or c.isspace())
        total_chars = len(text)
        if total_chars > 100 and readable_chars / total_chars < 0.3:
            return True

        return False

    def _parse_pdf_with_ocr(self, file_path: Path) -> Dict:
        """
        Parse PDF using OCR (Optical Character Recognition).

        This method converts PDF pages to images and uses Tesseract OCR
        to extract text. Used as fallback when pdfplumber fails.

        Args:
            file_path: Path to PDF file

        Returns:
            Dictionary with patient and blood_tests data
        """
        if not PYMUPDF_AVAILABLE:
            raise DataLoaderError(
                "PyMuPDF not available. Install it with: pip install PyMuPDF"
            )
        if not OCR_AVAILABLE:
            raise DataLoaderError(
                "OCR libraries not available. Install them with: pip install pytesseract Pillow"
            )

        try:
            doc = fitz.open(str(file_path))

            # Extract text from all pages using OCR
            full_text = ""
            for page_num, page in enumerate(doc):
                # Convert page to image
                mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better OCR
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                img = Image.open(io.BytesIO(img_data))

                # Perform OCR with Polish language
                page_text = pytesseract.image_to_string(
                    img, lang="pol", config="--psm 6"
                )
                full_text += page_text + "\n"

            doc.close()

            # Check if OCR produced usable text
            if not full_text.strip():
                raise DataLoaderError(
                    "OCR failed to extract text from PDF. "
                    "The PDF may be a scan-only document or use unsupported encoding."
                )

            # Parse the OCR'd text using existing text extraction method
            patient_data, blood_tests = self._parse_text_content(full_text)

            return {
                "patient": patient_data.to_dict(),
                "blood_tests": [test.to_dict() for test in blood_tests],
            }

        except Exception as e:
            raise DataLoaderError(f"Failed to parse PDF with OCR: {str(e)}")

    def _parse_text_content(self, text: str) -> tuple:
        """
        Parse patient data and blood tests from extracted text.

        Args:
            text: Extracted text content

        Returns:
            Tuple of (PatientData, List[BloodTest])
        """
        # Extract patient information using regex
        name_match = re.search(r"[Ii]mi[ęe][:\s]+([A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+)", text)
        surname_match = re.search(
            r"[Nn]azwisko[:\s]+([A-ZĄĆĘŁŃÓŚŹŻ][a-ząćęłńóśźż]+)", text
        )
        age_match = re.search(r"[Ww]iek[:\s]+(\d+)", text)

        patient_data = PatientData(
            name=name_match.group(1) if name_match else "",
            surname=surname_match.group(1) if surname_match else "",
            age=int(age_match.group(1)) if age_match else 0,
            conditions=[],
        )

        # Extract blood tests using multiple regex patterns
        blood_tests = []

        # Common blood test keywords to validate results
        valid_keywords = [
            r"witamina",
            r"wit\.",
            r"hemoglobina",
            r"cholesterol",
            r"glukoza",
            r"żelazo",
            r"ferrytyna",
            r"TSH",
            r"testosteron",
            r"kortyzol",
            r"trombocyty",
            r"leukocyty",
            r"erytrocyty",
            r"hematokryt",
            r"crp",
            r"mocznik",
            r"kreatynina",
            r"HDL",
            r"LDL",
            r"triglicerydy",
        ]

        # Pattern 1: "Test Name: Value Unit" (with colon)
        test_pattern_1 = r"([A-ZĄĆĘŁŃÓŚŹŻ][a-zA-Ząćęłńóśźż\s\.\-()]+?)[:\s]+([\d.,]+)\s*([a-zA-Z/%\*\^\s]*)"

        # Pattern 2: Table-like format with rows
        test_pattern_2 = r"([A-ZĄĆĘŁŃÓŚŹŻ][a-zA-Ząćęłńóśźż\s\.\-()]+?)\s+([\d.,]+)\s+([a-zA-Z/%\*\^]+)"

        for match in re.finditer(test_pattern_1, text):
            if self._add_blood_test_from_match(match, blood_tests, valid_keywords):
                continue

        for match in re.finditer(test_pattern_2, text):
            if self._add_blood_test_from_match(match, blood_tests, valid_keywords):
                continue

        return patient_data, blood_tests

    def _add_blood_test_from_match(
        self, match, blood_tests: List[BloodTest], valid_keywords: List[str] = None
    ) -> bool:
        """
        Add a blood test from regex match to the list.

        Args:
            match: Regex match object
            blood_tests: List to add the blood test to
            valid_keywords: List of keywords to validate blood test names

        Returns:
            True if blood test was added, False otherwise
        """
        try:
            name = match.group(1).strip()
            value_str = match.group(2).replace(",", ".")
            value = float(value_str)
            unit = match.group(3).strip() if len(match.groups()) > 2 else ""

            if not unit:
                unit = " "

            # Skip if name looks like header/footer text
            skip_patterns = [
                r"ul\.?",
                r"numer",
                r"nr\s",
                r"data",
                r"strona",
                r"oddział",
                r"kod",
                r"adres",
                r"tel\.?",
                r"fax",
                r"email",
                r"pacjent",
                r"księga",
                r"rejestrowy",
                r"podmiot",
                r"leczniczy",
            ]

            name_lower = name.lower()
            for pattern in skip_patterns:
                if re.search(pattern, name_lower):
                    return False

            # Skip if name is too long or contains only addresses/numbers
            if len(name) > 100 or len(name.split()) > 10:
                return False

            # If valid_keywords provided, check if name contains any
            if valid_keywords:
                if not any(
                    re.search(keyword, name_lower) for keyword in valid_keywords
                ):
                    # Still add if it looks like a medical test (short name, medical units)
                    if len(name.split()) > 3 or not any(
                        unit in unit
                        for unit in ["mg", "ng", "IU", "U", "%", "g/L", "mmol/L"]
                    ):
                        return False

            # Avoid duplicates
            if not any(bt.name == name for bt in blood_tests):
                blood_tests.append(BloodTest(name=name, value=value, unit=unit))
                return True

        except (ValueError, IndexError):
            pass

        return False

